from pathlib import Path
from typing import Any

import httpx
import pytest

from report_system.domain import Document, DocumentStatus, DocumentType, ValidationResult
from report_system.llm import (
    FakeLLMProvider,
    OllamaProvider,
    extract_document,
    generate_report,
    load_examples,
    validate_semantics,
)


ROOT = Path(__file__).parents[1]


def test_extraction_uses_schema_and_preserves_only_returned_facts() -> None:
    provider = FakeLLMProvider(
        [{"title": "Акт E-17", "sections": [{"name": "Процесс", "records": []}]}]
    )
    document = extract_document(
        provider,
        ROOT / "prompts",
        DocumentType.MANUFACTURING_ACT,
        "Сушили образец E-17.",
    )
    assert document.status == DocumentStatus.EXTRACTED
    assert document.sections[0].name == "Процесс"
    assert provider.requests[0]["json_schema"] is not None
    assert "pressure" not in document.model_dump_json()


def test_extraction_rejects_unsupported_type_and_blank_text() -> None:
    provider = FakeLLMProvider([])
    with pytest.raises(ValueError, match="not supported"):
        extract_document(provider, ROOT / "prompts", DocumentType.TEST_ACT, "данные")
    with pytest.raises(ValueError, match="cannot be empty"):
        extract_document(provider, ROOT / "prompts", DocumentType.TEST_PROTOCOL, "  ")


def test_generation_requires_confirmation() -> None:
    provider = FakeLLMProvider(["  Технический текст.  "])
    document = Document(document_type=DocumentType.TEST_PROTOCOL)
    with pytest.raises(ValueError, match="confirmed"):
        generate_report(provider, ROOT / "prompts", ROOT / "examples", document)
    document.status = DocumentStatus.CONFIRMED
    assert generate_report(provider, ROOT / "prompts", ROOT / "examples", document) == "Технический текст."


def test_semantic_validation_is_structured() -> None:
    provider = FakeLLMProvider([{"valid": False, "issues": [{"severity": "error", "statement": "печь", "reason": "не задана"}]}])
    result = validate_semantics(
        provider,
        ROOT / "prompts",
        Document(document_type=DocumentType.TEST_PROTOCOL),
        "Испытание в печи",
    )
    assert isinstance(result, ValidationResult)
    assert not result.valid


def test_ollama_is_loopback_only_and_disables_thinking(monkeypatch: pytest.MonkeyPatch) -> None:
    with pytest.raises(ValueError, match="loopback"):
        OllamaProvider("http://localhost.example:11434", "model")
    captured: dict[str, Any] = {}

    def fake_post(client: httpx.Client, url: str, *, json: dict[str, Any]) -> httpx.Response:
        captured.update(json)
        return httpx.Response(200, json={"response": '{"ok": true}'}, request=httpx.Request("POST", url))

    monkeypatch.setattr(httpx.Client, "post", fake_post)
    result = OllamaProvider("http://127.0.0.1:11434", "qwen3.5:9b").generate(
        "prompt", json_schema={"type": "object"}
    )
    assert result == {"ok": True}
    assert captured["think"] is False


def test_examples_load_from_txt_md_and_docx(tmp_path: Path) -> None:
    from docx import Document as DocxDocument

    directory = tmp_path / "protocols"
    directory.mkdir()
    (directory / "01.txt").write_text("TXT", encoding="utf-8")
    (directory / "02.md").write_text("MD", encoding="utf-8")
    docx = DocxDocument()
    docx.add_paragraph("DOCX")
    docx.save(directory / "03.docx")
    assert load_examples(tmp_path, DocumentType.TEST_PROTOCOL) == ["TXT", "MD", "DOCX"]

