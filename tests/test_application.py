from pathlib import Path

import pytest
from docx import Document as DocxDocument

from report_system.application import ReportApplication, build_test_act
from report_system.config import Settings
from report_system.docx import generate_docx
from report_system.domain import (
    Document,
    DocumentStatus,
    DocumentType,
    Parameter,
    Record,
    Section,
    SourceReference,
)
from report_system.llm import FakeLLMProvider
from report_system.validation import validate_facts


ROOT = Path(__file__).parents[1]


def settings(tmp_path: Path) -> Settings:
    return Settings(
        database_url=f"sqlite:///{tmp_path / 'app.db'}",
        output_dir=tmp_path / "output",
        prompts_dir=ROOT / "prompts",
        templates_dir=ROOT / "templates",
    )


def protocol(identifier: str, capacity: float, deviation: bool = False) -> Document:
    records = [
        Record(
            type="measurement",
            name="Ёмкость",
            parameters=[
                Parameter(
                    name="Значение",
                    value=capacity,
                    unit="А·ч",
                    source=SourceReference(source_type="user_input", raw_text_fragment=str(capacity)),
                )
            ],
        )
    ]
    if deviation:
        records.append(Record(type="deviation", name="Отклонение", parameters=[]))
    return Document(
        id=identifier,
        document_type=DocumentType.TEST_PROTOCOL,
        title=f"Протокол {identifier}",
        status=DocumentStatus.CONFIRMED,
        sections=[Section(name="Результаты", records=records)],
        conclusion="Испытание завершено",
    )


def test_deterministic_validation_detects_unknown_facts() -> None:
    document = Document(
        document_type=DocumentType.MANUFACTURING_ACT,
        title="Акт E-17",
        sections=[
            Section(
                name="Процесс",
                records=[
                    Record(
                        type="process",
                        name="Сушка",
                        parameters=[
                            Parameter(
                                name="Температура",
                                value=120,
                                source=SourceReference(
                                    source_type="user_input",
                                    raw_text_fragment="120",
                                ),
                            )
                        ],
                    )
                ],
            )
        ],
    )
    valid = validate_facts(document, "Образец E-17 сушили при 120 °C.")
    invalid = validate_facts(document, "Образец X-99 сушили при 130 °C.")
    assert valid.valid
    assert {issue.check for issue in invalid.issues} == {"unknown_number", "unknown_identifier"}


def test_deterministic_validation_ignores_section_numbering() -> None:
    document = Document(document_type=DocumentType.MANUFACTURING_ACT)
    result = validate_facts(document, "1. Общие сведения\n2) Результаты")
    assert result.valid


def test_test_act_is_built_deterministically_with_provenance() -> None:
    act = build_test_act([protocol("P-001", 4.81), protocol("P-002", 4.77, True)])
    counts = {item.name: item.value for item in act.sections[1].records[0].parameters}
    assert counts == {
        "Количество протоколов": 2,
        "Количество испытаний": 2,
        "Количество результатов": 2,
        "Количество отклонений": 1,
    }
    assert all(item.source.source_type == "deterministic_aggregation" for item in act.sections[1].records[0].parameters)
    results = next(section for section in act.sections if section.name == "Результаты испытаний")
    assert results.records[0].parameters[0].source.source_id == "P-001"


def test_test_act_rejects_wrong_or_unconfirmed_input() -> None:
    draft = protocol("P-001", 4.81)
    draft.status = DocumentStatus.EXTRACTED
    with pytest.raises(ValueError, match="confirmed"):
        build_test_act([draft])
    with pytest.raises(ValueError, match="not a test protocol"):
        build_test_act([Document(document_type=DocumentType.MANUFACTURING_ACT, status=DocumentStatus.CONFIRMED)])


def test_complete_generation_and_docx_flow(tmp_path: Path) -> None:
    provider = FakeLLMProvider(
        [
            {"title": "Акт E-17", "sections": []},
            "Получен образец E-17.",
            {"valid": True, "issues": []},
        ]
    )
    application = ReportApplication(settings(tmp_path), provider)
    extracted = application.extract(DocumentType.MANUFACTURING_ACT, "Получен образец E-17.")
    assert application.repository.get(extracted.id) is None
    confirmed = application.confirm(extracted)
    document, validation, output = application.generate(confirmed.id)
    assert document.status == DocumentStatus.GENERATED
    assert validation.valid and output and output.exists()
    assert application.repository.get(document.id).revision == 1  # type: ignore[union-attr]


def test_invalid_report_is_revised_before_export(tmp_path: Path) -> None:
    application = ReportApplication(
        settings(tmp_path),
        FakeLLMProvider(["Добавлено значение 999.", "Испытание выполнено."]),
    )
    source = application.confirm(Document(document_type=DocumentType.TEST_PROTOCOL))
    document, validation, output = application.generate(source.id, semantic_validation=False)
    assert validation.valid
    assert output and output.exists()
    assert document.status == DocumentStatus.GENERATED
    assert document.generated_text == "Испытание выполнено."
    assert "999" not in document.generated_text


def test_report_is_not_stored_or_exported_when_revision_fails(tmp_path: Path) -> None:
    application = ReportApplication(
        settings(tmp_path),
        FakeLLMProvider(["Добавлено значение 999.", "Добавлено значение 888."]),
    )
    source = application.confirm(Document(document_type=DocumentType.TEST_PROTOCOL))
    document, validation, output = application.generate(source.id, semantic_validation=False)
    stored = application.repository.get(document.id)
    assert not validation.valid
    assert output is None
    assert document.status == DocumentStatus.CONFIRMED
    assert document.generated_text is None
    assert stored is not None and stored.generated_text is None


def test_semantic_hallucination_is_removed_and_checked_again(tmp_path: Path) -> None:
    provider = FakeLLMProvider(
        [
            "Получен образец E-17 в печи.",
            {
                "valid": False,
                "issues": [
                    {
                        "severity": "error",
                        "statement": "в печи",
                        "reason": "Оборудование не подтверждено",
                    }
                ],
            },
            "Получен образец E-17.",
            {"valid": True, "issues": []},
        ]
    )
    application = ReportApplication(settings(tmp_path), provider)
    source = application.confirm(
        Document(
            document_type=DocumentType.MANUFACTURING_ACT,
            title="Образец E-17",
        )
    )
    document, validation, output = application.generate(source.id)
    assert validation.valid
    assert output and output.exists()
    assert document.generated_text == "Получен образец E-17."
    assert "печи" not in document.generated_text
    assert len(provider.requests) == 4


def test_real_templates_are_usable(tmp_path: Path) -> None:
    for name in ("manufacturing_act.docx", "test_protocol.docx", "test_act.docx"):
        text = "\n".join(paragraph.text for paragraph in DocxDocument(ROOT / "templates" / name).paragraphs)
        assert "{{title}}" in text and "{{document_id}}" in text

    document = Document(
        id="MA-CHECK-001",
        document_type=DocumentType.MANUFACTURING_ACT,
        title="Проверочный акт",
        status=DocumentStatus.GENERATED,
        generated_text="Проверочный текст.",
        sections=[
            Section(
                name="Параметры",
                records=[
                    Record(
                        type="measurement",
                        name="Температура",
                        parameters=[
                            Parameter(
                                name="Значение",
                                value=120,
                                unit="°C",
                                source=SourceReference(
                                    source_type="user_input",
                                    raw_text_fragment="120 °C",
                                ),
                            )
                        ],
                    )
                ],
            )
        ],
    )
    output = generate_docx(document, ROOT / "templates", tmp_path)
    generated_docx = DocxDocument(output)
    paragraphs = [paragraph.text for paragraph in generated_docx.paragraphs]
    text = "\n".join(paragraphs)
    assert text.count("Проверочный акт") == 1
    assert "MA-CHECK-001" in text
    assert len(generated_docx.tables) == 1
    assert "<w:tblBorders>" in generated_docx.tables[0]._tbl.xml
