import json
import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from report_system.domain import Document, DocumentType, Parameter


def generate_docx(document: Document, templates_dir: Path, output_dir: Path) -> Path:
    if not document.generated_text:
        raise ValueError("generated text is required for DOCX")
    template_name = {
        DocumentType.MANUFACTURING_ACT: "manufacturing_act.docx",
        DocumentType.TEST_PROTOCOL: "test_protocol.docx",
        DocumentType.TEST_ACT: "test_act.docx",
    }[document.document_type]
    template = templates_dir / template_name
    docx = DocxDocument(template) if template.exists() else DocxDocument()
    replacements = {
        "{{title}}": document.title or "",
        "{{document_id}}": document.id,
        "{{generated_text}}": document.generated_text,
    }
    had_title = False
    for paragraph in docx.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                had_title = had_title or placeholder == "{{title}}"
                paragraph.text = paragraph.text.replace(placeholder, value)
    if not had_title:
        title = docx.add_heading(document.title or _default_title(document.document_type), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx.add_heading("Содержание", level=1)
    for block in document.generated_text.split("\n"):
        if block.strip():
            docx.add_paragraph(block.strip())
    docx.add_heading("Структурированные данные", level=1)
    for section in document.sections:
        docx.add_heading(section.name, level=2)
        for record in section.records:
            docx.add_heading(record.name, level=3)
            if not record.parameters:
                continue
            table = docx.add_table(rows=1, cols=3)
            _apply_table_grid(docx, table)
            for cell, label in zip(table.rows[0].cells, ("Параметр", "Значение", "Источник"), strict=True):
                cell.text = label
            for parameter in record.parameters:
                cells = table.add_row().cells
                cells[0].text = parameter.name
                cells[1].text = _parameter_text(parameter)
                cells[2].text = _source_text(parameter)

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_id = re.sub(r"[^A-Za-zА-Яа-яЁё0-9_.-]+", "_", document.id).strip("._") or "document"
    output = output_dir / f"{document.document_type.value}_{safe_id}.docx"
    docx.save(output)
    return output


def _parameter_text(parameter: Parameter) -> str:
    value = json.dumps(parameter.value, ensure_ascii=False) if isinstance(parameter.value, (dict, list)) else parameter.value
    return f"{value} {parameter.unit or ''}".strip()


def _source_text(parameter: Parameter) -> str:
    return parameter.source.source_id or parameter.source.raw_text_fragment or parameter.source.source_path or ""


def _apply_table_grid(document: DocxDocumentType, table: Table) -> None:
    grid_style = next(
        (
            style
            for style in document.styles
            if style.type == WD_STYLE_TYPE.TABLE and style.style_id == "TableGrid"
        ),
        None,
    )
    if grid_style is not None:
        table.style = grid_style
        return

    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "auto")


def _default_title(document_type: DocumentType) -> str:
    return {
        DocumentType.MANUFACTURING_ACT: "Акт наработки",
        DocumentType.TEST_PROTOCOL: "Протокол испытаний",
        DocumentType.TEST_ACT: "Акт испытаний",
    }[document_type]
