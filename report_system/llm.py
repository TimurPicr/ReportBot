from collections import deque
from collections.abc import Mapping
from pathlib import Path
from typing import Any, Protocol
from urllib.parse import urlparse
import json

import httpx

from report_system.domain import (
    Document,
    DocumentStatus,
    DocumentType,
    Section,
    ValidationResult,
)


class LLMProvider(Protocol):
    def generate(
        self,
        prompt: str,
        *,
        json_schema: Mapping[str, Any] | None = None,
        temperature: float = 0.0,
    ) -> str | dict[str, Any]: ...


class OllamaProvider:
    def __init__(self, base_url: str, model: str, timeout: float = 600.0) -> None:
        parsed = urlparse(base_url)
        if parsed.scheme != "http" or parsed.hostname not in {"127.0.0.1", "localhost", "::1"}:
            raise ValueError("Ollama provider only permits loopback URLs")
        self.base_url = base_url.rstrip("/")
        self.model = model
        self.timeout = timeout

    def generate(
        self,
        prompt: str,
        *,
        json_schema: Mapping[str, Any] | None = None,
        temperature: float = 0.0,
    ) -> str | dict[str, Any]:
        payload: dict[str, Any] = {
            "model": self.model,
            "prompt": prompt,
            "stream": False,
            "think": False,
            "options": {"temperature": temperature},
        }
        if json_schema is not None:
            payload["format"] = dict(json_schema)
        with httpx.Client(trust_env=False, timeout=self.timeout) as client:
            response = client.post(f"{self.base_url}/api/generate", json=payload)
        response.raise_for_status()
        body = response.json()
        text = body.get("response", "")
        if not text.strip():
            raise RuntimeError(
                "Ollama returned an empty response "
                f"(done_reason={body.get('done_reason')}, thinking_length={len(body.get('thinking', ''))})"
            )
        if json_schema is None:
            return text
        try:
            result = json.loads(text)
        except json.JSONDecodeError as error:
            raise ValueError(f"Ollama returned invalid JSON: {text[:200]!r}") from error
        if not isinstance(result, dict):
            raise ValueError("structured Ollama response must be a JSON object")
        return result


class FakeLLMProvider:
    def __init__(self, responses: list[str | dict[str, Any]]) -> None:
        self.responses = deque(responses)
        self.requests: list[dict[str, Any]] = []

    def generate(
        self,
        prompt: str,
        *,
        json_schema: Mapping[str, Any] | None = None,
        temperature: float = 0.0,
    ) -> str | dict[str, Any]:
        self.requests.append({"prompt": prompt, "json_schema": json_schema, "temperature": temperature})
        if not self.responses:
            raise RuntimeError("FakeLLMProvider has no responses left")
        return self.responses.popleft()


def _prompt(prompts_dir: Path, filename: str) -> str:
    return (prompts_dir / filename).read_text(encoding="utf-8")


def extract_document(
    provider: LLMProvider,
    prompts_dir: Path,
    document_type: DocumentType,
    raw_input: str,
) -> Document:
    prompt_files = {
        DocumentType.MANUFACTURING_ACT: "extract_manufacturing.txt",
        DocumentType.TEST_PROTOCOL: "extract_protocol.txt",
    }
    if document_type not in prompt_files:
        raise ValueError(f"extraction is not supported for {document_type}")
    if not raw_input.strip():
        raise ValueError("raw input cannot be empty")

    full_schema = Document.model_json_schema()
    fields = ("title", "metadata", "sections", "conclusion")
    schema = {
        "$defs": full_schema.get("$defs", {}),
        "type": "object",
        "additionalProperties": False,
        "properties": {name: full_schema["properties"][name] for name in fields},
    }
    request = (
        f"{_prompt(prompts_dir, prompt_files[document_type])}\n\n"
        f"Тип документа: {document_type.value}\nИсходный текст пользователя:\n---\n{raw_input}\n---"
    )
    response = provider.generate(request, json_schema=schema, temperature=0.0)
    data = json.loads(response) if isinstance(response, str) else response
    unknown = set(data) - set(fields)
    if unknown:
        raise ValueError(f"unexpected extraction fields: {sorted(unknown)}")
    return Document(
        document_type=document_type,
        title=data.get("title"),
        metadata=data.get("metadata", {}),
        sections=[Section.model_validate(item) for item in data.get("sections", [])],
        conclusion=data.get("conclusion"),
        raw_input=raw_input,
        status=DocumentStatus.EXTRACTED,
    )


def load_examples(root: Path, document_type: DocumentType, limit: int = 3) -> list[str]:
    from docx import Document as DocxDocument

    directory = root / {
        DocumentType.MANUFACTURING_ACT: "manufacturing",
        DocumentType.TEST_PROTOCOL: "protocols",
        DocumentType.TEST_ACT: "test_acts",
    }[document_type]
    if not directory.exists() or limit <= 0:
        return []
    result: list[str] = []
    for path in sorted(directory.iterdir()):
        if path.suffix.lower() in {".txt", ".md"}:
            result.append(path.read_text(encoding="utf-8"))
        elif path.suffix.lower() == ".docx":
            document = DocxDocument(path)
            result.append("\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text))
        if len(result) == limit:
            break
    return result


def generate_report(provider: LLMProvider, prompts_dir: Path, examples_dir: Path, document: Document) -> str:
    if document.status != DocumentStatus.CONFIRMED:
        raise ValueError("document must be confirmed before generation")
    prompt_file = {
        DocumentType.MANUFACTURING_ACT: "generate_manufacturing.txt",
        DocumentType.TEST_PROTOCOL: "generate_protocol.txt",
        DocumentType.TEST_ACT: "generate_test_act.txt",
    }[document.document_type]
    facts = {
        "document_type": document.document_type.value,
        "title": document.title,
        "metadata": document.metadata,
        "sections": [section.model_dump(mode="json") for section in document.sections],
        "source_document_ids": document.source_document_ids,
        "conclusion": document.conclusion,
    }
    request = (
        f"{_prompt(prompts_dir, prompt_file)}\n\n"
        "ПРИМЕРЫ СТИЛЯ (не являются источником фактов):\n"
        f"{json.dumps([item[:6000] for item in load_examples(examples_dir, document.document_type)], ensure_ascii=False)}\n\n"
        f"ПОДТВЕРЖДЁННЫЕ ФАКТЫ:\n{json.dumps(facts, ensure_ascii=False)}"
    )
    response = provider.generate(request, temperature=0.1)
    if not isinstance(response, str):
        raise TypeError("report generator expects a text response")
    return response.strip()


def validate_semantics(
    provider: LLMProvider,
    prompts_dir: Path,
    document: Document,
    generated_text: str,
) -> ValidationResult:
    request = (
        f"{_prompt(prompts_dir, 'validate_report.txt')}\n\n"
        f"STRUCTURED DATA:\n{document.model_dump_json(exclude={'raw_input', 'generated_text'})}\n\n"
        f"REPORT TEXT:\n{generated_text}"
    )
    response = provider.generate(request, json_schema=ValidationResult.model_json_schema(), temperature=0.0)
    return ValidationResult.model_validate_json(response) if isinstance(response, str) else ValidationResult.model_validate(response)
